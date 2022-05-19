import json
from docx import Document

import re
from docx.shared import Pt,RGBColor, Inches


doc1=Document()


import requests

cookies = {
    'Hm_lvt_913838bb97e4853192b599c0f51faf2b': '1646565056,1646722049,1646724520,1646730696',
    'isReadeNotice': 'true',
    'Hm_lvt_cff27b63a4732c7d33b270aebc954184': '1651573215,1651586342,1651650473,1651659169',
    'Hm_lpvt_cff27b63a4732c7d33b270aebc954184': '1651659169',
}

headers = {
    'Accept': 'application/json, text/plain, */*',
    'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8,zh-SA;q=0.7,en-SA;q=0.6',
    'Connection': 'keep-alive',
    'Content-Type': 'application/json;charset=UTF-8',
    # Requests sorts cookies= alphabetically
    # 'Cookie': 'Hm_lvt_913838bb97e4853192b599c0f51faf2b=1646565056,1646722049,1646724520,1646730696; isReadeNotice=true; Hm_lvt_cff27b63a4732c7d33b270aebc954184=1651573215,1651586342,1651650473,1651659169; Hm_lpvt_cff27b63a4732c7d33b270aebc954184=1651659169',
    'Origin': 'http://tp.1safety.cc',
    'P-Token': 'c42f6264f3f846938240f40d248d1f99',
    'Referer': 'http://tp.1safety.cc/welcome',
    'User-Agent': 'Mozilla/5.0 (Linux; Android 6.0; Nexus 5 Build/MRA58N) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Mobile Safari/537.36',
    'appType': '0',
    'userId': '0sbbh11we9tx',
}

json_data = {
    'user_id': '0sbbh11we9tx',
    'project_id': '1najd51adavd',
    'page_size': 1000,
    'page_index': 1,
}

response = requests.post('http://tp.1safety.cc/tp-api/v1.0/exercise/page', cookies=cookies, headers=headers, json=json_data, verify=False)







def save_data(response):
    data_json = json.loads(response.content.decode())
    # print(data_json)
    question_list = data_json['data']['datas']
    for question in range(len(question_list)):
        question_index = str(question + 1)
        # print(question_index)
        if question_list[question]["chr_type"] == "01":
            question_type = "(单选题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "02":
            question_type = "(多选题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "03":
            question_type = "(判断题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "04":
            question_type = "(填空题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "05":
            question_type = "(简答题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "06":
            question_type = "(论述题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "07":
            question_type = "(分析题)、"
            # print(question_type)
        if question_list[question]["chr_type"] == "08":
            question_type = "(绘图题)、"
            # print(question_type)

        if "var_questions_content" in question_list[question]:
            question_content = question_list[question]["var_questions_content"]

            doc1.add_paragraph(question_index + question_type + question_content)

            question_title_files = question_list[question]["question_title_files"]


            for question_title_file in range(len(question_title_files)):
                if "mime_url" in question_title_files[question_title_file]:
                    question_title_img = question_title_files[question_title_file]["mime_url"]
                    opton_title_pic_name = re.search(r'srcFid=(.*)&expires?', question_title_img).group(1)
                    #print('http://cdns001.bosafe.com/api/v1/file/view/' + opton_title_pic_name)
                    with open('images/' + opton_title_pic_name, 'wb') as f:
                        f.write(
                            requests.get('http://cdns001.bosafe.com/api/v1/file/view/' + opton_title_pic_name).content)
                    doc1.add_picture('images/' + opton_title_pic_name, width=Inches(2))


            # print(question_content)
            if "options" in question_list[question]:
                question_options = question_list[question]["options"]
                for opt in range(len(question_options)):
                    opt_type = question_options[opt]["o_type"]
                    opt_content = question_options[opt]["option_content"]
                    doc1.add_paragraph(opt_type + "、" + opt_content)
                    if "mime_url" in question_options[opt]:
                        question_option_img = question_options[opt]["mime_url"]
                        opton_pic_name = re.search(r'srcFid=(.*)&expires?', question_option_img).group(1)
                        # print(pic_id)
                        with open('images/' + opton_pic_name, 'wb') as f:
                            f.write(
                                requests.get('http://cdns001.bosafe.com/api/v1/file/view/' + opton_pic_name).content)
                        doc1.add_picture('images/' + opton_pic_name, width=Inches(2))

        if "var_answer" in question_list[question]:
            question_var_answer = question_list[question]["var_answer"]
            doc1.add_paragraph("正确答案：" + question_var_answer)
            # print("正确答案："+question_var_answer)

        if "answer_desc" in question_list[question]:
            question_answer_desc = question_list[question]["answer_desc"]
            for item in range(len(question_answer_desc)):
                question_answer_content = question_answer_desc[item]['answer_content']
                question_answer_files = question_answer_desc[item]['answer_files']
                doc1.add_paragraph("正确答案：" + question_answer_content)

                for file in question_answer_files:
                    if "mime_url" in file:
                        question_answer_img = file["mime_url"]
                        pic_name = re.search(r'srcFid=(.*)&expires?', question_answer_img).group(1)
                        # print(pic_id)
                        with open('images/' + pic_name, 'wb') as f:
                            f.write(requests.get('http://cdns001.bosafe.com/api/v1/file/view/' + pic_name).content)
                        doc1.add_picture('images/' + pic_name, width=Inches(6))
                        # print('http://cdns001.bosafe.com/api/v1/file/view/'+pic_name)

        doc1.save('answer.docx')
save_data(response)


